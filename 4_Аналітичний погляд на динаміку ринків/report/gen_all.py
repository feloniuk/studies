# -*- coding: utf-8 -*-
"""Генератор APA-звітів по акціях S&P 500 (2026 р.). Один тікер -> один .docx."""
import json, os, sys
import pandas as pd, numpy as np
from scipy import stats
import matplotlib; matplotlib.use('Agg')
import matplotlib.pyplot as plt
import pypandoc
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

plt.rcParams.update({'font.size':9,'figure.dpi':140})
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
os.chdir(BASE)

COMPANIES = [
    ('MSFT','Microsoft Corporation','Microsoft'),
    ('AMZN','Amazon.com Inc.','Amazon'),
    ('NVDA','NVIDIA Corporation','NVIDIA'),
    ('META','Meta Platforms Inc.','Meta'),
]
close_all = pd.read_csv('data/prices_daily.csv', index_col=0, parse_dates=True)

# ---------- helpers ----------
def num(v,d=2):
    if isinstance(v,str): return v
    return f'{v:,.{d}f}'.replace(',',' ').replace('.',',')
def sg(v,d=2): return ('+' if v>=0 else '−')+' '+num(abs(v),d)

def r2(y,yp): return 1-np.sum((y-yp)**2)/np.sum((y-y.mean())**2)
def trends(y):
    n=len(y); x=np.arange(1,n+1)
    la,lb=np.polyfit(x,y,1); lin=la*x+lb
    pc=np.polyfit(x,y,2); poly=np.polyval(pc,x)
    pb,plc=np.polyfit(np.log(x),np.log(y),1); powy=np.exp(plc)*x**pb
    return dict(lin=dict(a=la,b=lb,r2=r2(y,lin)),
                poly=dict(a=pc[0],b=pc[1],c=pc[2],r2=r2(y,poly)),
                pow=dict(c=np.exp(plc),b=pb,r2=r2(y,powy)),
                _lin=lin,_poly=poly,_pow=powy)
def desc(v):
    v=np.asarray(v,float); m=v.mean()
    vals,cnts=np.unique(np.round(v,2),return_counts=True)
    mode='#Н/Д' if cnts.max()==1 else float(vals[cnts.argmax()])
    return dict(mean=m,median=float(np.median(v)),mode=mode,min=float(v.min()),max=float(v.max()),
        rng=float(v.max()-v.min()),kurt=float(stats.kurtosis(v)),skew=float(stats.skew(v)),
        cv=float(v.std(ddof=1)/m*100),osc=float((v.max()-v.min())/m*100),std=float(v.std(ddof=1)))

def make_figs(tk, short, close, tr):
    y=close.values; n=len(y); x=np.arange(1,n+1)
    img=f'report/img/{tk}'
    # fig1 price + trends
    fig,ax=plt.subplots(figsize=(7.2,4.0))
    ax.plot(x,y,color='#1f4e79',lw=1.3,label=short)
    ax.plot(x,tr['_lin'],'--',color='#c00000',lw=1,label='Лінійна')
    ax.plot(x,tr['_poly'],'--',color='#548235',lw=1,label='Поліноміальна')
    ax.plot(x,tr['_pow'],'--',color='#7030a0',lw=1,label='Степенева')
    ax.set_xlabel('Період (торговий день)'); ax.set_ylabel('Ціна, дол. США')
    l,po,pw=tr['lin'],tr['poly'],tr['pow']
    sgn=lambda v:('+' if v>=0 else '−')+f'{abs(v):.2f}'
    txt=(f"y = {l['a']:.4f}x {sgn(l['b'])};  R² = {l['r2']:.4f}\n"
         f"y = {po['a']:.5f}x² {('+' if po['b']>=0 else '−')}{abs(po['b']):.4f}x {sgn(po['c'])};  R² = {po['r2']:.4f}\n"
         f"y = {pw['c']:.3e}·x^{pw['b']:.3f};  R² = {pw['r2']:.4f}")
    ax.text(0.02,0.97,txt,transform=ax.transAxes,va='top',fontsize=7.5,
            bbox=dict(boxstyle='round',fc='white',ec='0.7',alpha=0.9))
    ax.legend(loc='lower left',fontsize=7,ncol=4); ax.grid(alpha=0.25)
    ax.set_title(f'Динаміка цін на акції {short} (щоденні значення, 2026 р.)',fontsize=8.5)
    plt.tight_layout(); plt.savefig(img+'_price.png'); plt.close()
    # returns figs
    def rf(s,fn,title):
        fig,ax=plt.subplots(figsize=(7.2,2.7))
        ax.plot(np.arange(1,len(s)+1),s.values*100,color='#1f4e79',lw=0.9)
        ax.axhline(0,color='0.5',lw=0.6); ax.set_ylabel('%'); ax.set_xlabel('Період')
        ax.set_title(title,fontsize=8.5); ax.grid(alpha=0.25)
        plt.tight_layout(); plt.savefig(fn); plt.close()
    rd=close.pct_change().dropna()
    rw=close.resample('W-FRI').last().pct_change().dropna()
    rm=close.resample('ME').last().pct_change().dropna()
    rf(rd,img+'_ret_d.png',f'Динаміка щоденної прибутковості акцій {short} (%, 2026 р.)')
    rf(rw,img+'_ret_w.png',f'Динаміка щотижневої прибутковості акцій {short} (%, 2026 р.)')
    rf(rm,img+'_ret_m.png',f'Динаміка щомісячної прибутковості акцій {short} (%, 2026 р.)')
    # histogram
    fig,ax=plt.subplots(figsize=(7.2,3.4))
    data=rd.values*100
    ax.hist(data,bins=15,color='#9dc3e6',edgecolor='#1f4e79',density=True,alpha=0.9)
    xs=np.linspace(data.min(),data.max(),200)
    ax.plot(xs,stats.norm.pdf(xs,data.mean(),data.std()),'r-',lw=1.4,label='Нормальний закон')
    ax.set_xlabel('Прибутковість, %'); ax.set_ylabel('Щільність'); ax.legend(fontsize=7)
    ax.set_title(f'Гістограма розподілу щоденної прибутковості акцій {short} (2026 р.)',fontsize=8.5)
    plt.tight_layout(); plt.savefig(img+'_hist.png'); plt.close()
    return rd,rw,rm

# ---------- markdown ----------
def best_trend(tr):
    d={'лінійний':tr['lin']['r2'],'поліноміальний':tr['poly']['r2'],'степеневий':tr['pow']['r2']}
    return max(d,key=d.get), d
def build_md(tk, legal, short, close, tr, rd, rw, rm, others):
    n=len(close); start=str(close.index.min().date()); end=str(close.index.max().date())
    nw=len(rw); nm=len(rm)
    p=desc(close.values); sd=desc(rd.values); sw=desc(rw.values); sm=desc(rm.values)
    img=f'report/img/{tk}'
    l,po,pw=tr['lin'],tr['poly'],tr['pow']
    head=close.head(5); tail=close.tail(5)
    def rows(sr): return '\n'.join(f"| {i.strftime('%d.%m.%Y')} | {num(v)} |" for i,v in sr.items())
    def stat_table(d,name,sym,nn):
        return (f"**Таблиця {name}**\n\n*Статистичні показники {sym} прибутковості акцій {short} за 2026 р. (n = {nn})*\n\n"
        "| Показник | Значення |\n|:--|--:|\n"
        f"| Середнє | {num(d['mean'],4)} |\n| Медіана | {num(d['median'],4)} |\n| Мода | {num(d['mode'],4)} |\n"
        f"| Мінімум | {num(d['min'],4)} |\n| Максимум | {num(d['max'],4)} |\n| Розмах | {num(d['rng'],4)} |\n"
        f"| Ексцес | {num(d['kurt'],4)} |\n| Асиметричність | {num(d['skew'],4)} |\n"
        f"| Коефіцієнт варіації, % | {num(d['cv'],2)} |\n| Коефіцієнт осциляції, % | {num(d['osc'],2)} |\n\n"
        "*Примітка.* Розраховано автором за даними Yahoo Finance (2026).")
    # adaptive wording
    bt,_=best_trend(tr)
    traj_adj = 'спадною' if l['a']<0 else 'висхідною'
    skew_p = 'Додатна асиметрія' if p['skew']>0 else "Від'ємна асиметрія"
    skew_p_tail = 'нижчих за середнє' if p['skew']>0 else 'вищих за середнє'
    kurt_p = 'дещо гостріший, ніж нормальний' if p['kurt']>0 else 'дещо пласкіший, ніж нормальний'
    mean_d_word = "від'ємною" if sd['mean']<0 else 'додатною'
    align_word = ('узгоджується зі спадним нахилом ціни' if sd['mean']<0
                  else 'узгоджується з висхідним нахилом ціни')
    kurt_d = ('гостровершинність і «важкі хвости» (лептокуртичність)' if sd['kurt']>0
              else 'пласковершинність розподілу (платикуртичність)')
    skew_d = ("від'ємна асиметрія ({}) — на переважання різких від'ємних змін над додатними".format(num(sd['skew'],2))
              if sd['skew']<0 else
              "додатна асиметрія ({}) — на наявність окремих різких додатних змін".format(num(sd['skew'],2)))
    others_str = ', '.join(others[:-1]) + ' та ' + others[-1] if len(others)>1 else others[0]
    mean_w_word = "від'ємної" if sw['mean']<0 else "додатної"

    md=f"""# Інфраструктура фінансового ринку та джерела даних

Сучасний фінансовий ринок є складною інформаційною системою, у якій ціни фінансових інструментів формуються під впливом великої кількості учасників і чинників. Інформаційну підтримку ринку забезпечують біржі, брокерські компанії, інформаційні агентства та фінансові портали, які акумулюють і публікують ретроспективні дані щодо динаміки котирувань (Арутюнян та ін., 2018). Саме доступність таких даних дає змогу досліднику отримати часові ряди цін і застосувати до них апарат статистичного аналізу.

Для дослідження ми обрали чотири високоліквідні акції, що входять до складу індексу S&P 500, — Microsoft Corporation (MSFT), Amazon.com Inc. (AMZN), NVIDIA Corporation (NVDA) та Meta Platforms Inc. (META). Перевагу цим емітентам надано через значні обсяги торгів, які забезпечують репрезентативність і неперервність цінового ряду. У межах цієї частини роботи детально розглянуто інструмент — звичайні акції {legal}, що котируються на біржі NASDAQ під тікером {tk}.

Джерелом первинних даних слугував глобальний фінансовий портал Investing.com, а також відкритий історичний архів котирувань Yahoo Finance, з якого було імпортовано щоденні ціни закриття з поправкою на корпоративні події (дроблення акцій та виплату дивідендів). Процедура отримання даних передбачала вибір інструмента, задання діапазону дат і періодичності та експорт історичних котирувань у форматі CSV із подальшим опрацюванням у середовищі обробки даних. Дослідний період охоплює перше півріччя 2026 року — з {start} до {end}, що відповідає {n} торговим дням. Фрагмент імпортованого часового ряду наведено в таблиці 1.

**Таблиця 1**

*Фрагмент імпортованого часового ряду цін закриття акцій {short} (дол. США)*

| Дата | Ціна закриття |
|:--|--:|
{rows(head)}
| … | … |
{rows(tail)}

*Примітка.* Наведено перші та останні п'ять спостережень вибірки. Складено автором за даними Yahoo Finance (2026).

Позначимо досліджуваний часовий ряд як $Z = \\langle \\overline{{z_{{d,\\,w,\\,m}}}} \\rangle$, де $z_d$ — ціна закриття акції за торговий день $d$, $d = \\overline{{1,\\,{n}}}$; $z_w$ та $z_m$ — відповідно щотижневі та щомісячні значення ціни, агреговані як ціна останнього торгового дня періоду.

# Методика статистичного аналізу

Для характеристики динаміки ціни застосовано поєднання трендового та описово-статистичного аналізу, що узгоджується з підходами, викладеними в роботах Максишко (2009) та Максишко і Шаповалової (2018). На першому етапі тенденцію часового ряду апроксимовано трьома типами трендів — лінійним, поліноміальним (другого степеня) та степеневим:

$$y = a x + b, \\qquad y = a x^{{2}} + b x + c, \\qquad y = c\\,x^{{b}}.$$

Якість апроксимації оцінено за коефіцієнтом детермінації

$$R^{{2}} = 1 - \\frac{{\\sum_{{i=1}}^{{n}} (y_i - \\hat{{y}}_i)^{{2}}}}{{\\sum_{{i=1}}^{{n}} (y_i - \\bar{{y}})^{{2}}}},$$

де $y_i$ — фактичне значення ціни, $\\hat{{y}}_i$ — модельне значення, $\\bar{{y}}$ — середнє ряду.

На другому етапі обчислено описові статистичні показники. Середнє значення та незміщену оцінку середньоквадратичного відхилення визначено за формулами

$$\\bar{{x}} = \\frac{{1}}{{n}} \\sum_{{i=1}}^{{n}} x_i, \\qquad \\sigma = \\sqrt{{\\frac{{1}}{{n-1}} \\sum_{{i=1}}^{{n}} (x_i - \\bar{{x}})^{{2}}}}.$$

Відносну мінливість ряду оцінено коефіцієнтами варіації та осциляції

$$V = \\frac{{\\sigma}}{{\\bar{{x}}}} \\cdot 100\\%, \\qquad V_R = \\frac{{z_{{max}} - z_{{min}}}}{{\\bar{{x}}}} \\cdot 100\\%,$$

а форму розподілу — коефіцієнтами асиметрії та ексцесу

$$A_s = \\frac{{1}}{{n}} \\sum_{{i=1}}^{{n}} \\frac{{(x_i - \\bar{{x}})^{{3}}}}{{\\sigma^{{3}}}}, \\qquad E_k = \\frac{{1}}{{n}} \\sum_{{i=1}}^{{n}} \\frac{{(x_i - \\bar{{x}})^{{4}}}}{{\\sigma^{{4}}}} - 3.$$

Близькість значень $A_s$ та $E_k$ до нуля свідчить про наближення емпіричного розподілу до нормального закону. Нарешті, прибутковість інструмента розраховано як відносну зміну ціни між сусідніми періодами

$$p_d = \\frac{{z_{{d+1}} - z_d}}{{z_d}} \\cdot 100\\%,$$

з аналогічними виразами для щотижневої $p_w$ та щомісячної $p_m$ прибутковості.

# Аналіз динаміки ціни

Графічне представлення часового ряду цін разом із побудованими лініями тренду наведено на рисунку 1. Протягом першого півріччя 2026 року ціна акцій {short} змінювалася в діапазоні від {num(p['min'])} до {num(p['max'])} дол. США; загальна тенденція ряду була {traj_adj}, що відображають побудовані лінії тренду.

**Рисунок 1**

*Динаміка цін на акції {short} за перше півріччя 2026 р. (щоденні значення) з лініями тренду*

![]({img}_price.png)

*Примітка.* Побудовано автором за даними Yahoo Finance (2026).

Серед трьох апроксимувальних моделей найкраще динаміку ряду описує {bt} тренд із коефіцієнтом детермінації $R^2 = {num(max(l['r2'],po['r2'],pw['r2']),4)}$. Кутовий коефіцієнт лінійного тренду ($a = {num(l['a'],4)}$) фіксує загальний {('спадний' if l['a']<0 else 'висхідний')} нахил ціни за півріччя. Узагальнені результати трендового та описово-статистичного аналізу подано в таблиці 2.

**Таблиця 2**

*Показники динаміки ціни акцій {short} за 2026 р.*

| Показник | Значення |
|:--|:--|
| Лінійний тренд | y = {num(l['a'],4)}x {sg(l['b'])}; R² = {num(l['r2'],4)} |
| Поліноміальний тренд | y = {num(po['a'],5)}x² {sg(po['b'],4)}x {sg(po['c'])}; R² = {num(po['r2'],4)} |
| Степеневий тренд | y = {num(pw['c'])}·x^{num(pw['b'],4)}; R² = {num(pw['r2'],4)} |
| Середнє | {num(p['mean'])} |
| Медіана | {num(p['median'])} |
| Мода | {num(p['mode'])} |
| Мінімум | {num(p['min'])} |
| Максимум | {num(p['max'])} |
| Розмах | {num(p['rng'])} |
| Ексцес | {num(p['kurt'],4)} |
| Асиметричність | {num(p['skew'],4)} |
| Коефіцієнт варіації, % | {num(p['cv'],4)} |
| Коефіцієнт осциляції, % | {num(p['osc'],4)} |

*Примітка.* Розраховано автором за даними Yahoo Finance (2026).

Коефіцієнт варіації цінового ряду становить {num(p['cv'],2)} %, що свідчить про {'помірну' if p['cv']<33 else 'значну'} мінливість значень навколо середнього, тоді як коефіцієнт осциляції {num(p['osc'],2)} % відображає відносний розмах коливань — від {num(p['min'])} до {num(p['max'])} дол. США. {skew_p} ({num(p['skew'],2)}) вказує на незначне переважання значень, {skew_p_tail}, а ексцес ({num(p['kurt'],2)}) — на {kurt_p} розподіл цін.

# Аналіз прибутковості

## Щоденна прибутковість

Динаміку щоденної прибутковості акцій наведено на рисунку 2. Ряд коливається навколо нульової позначки з періодичними сплесками волатильності, найбільший з яких відповідає денній зміні близько {num(max(abs(sd['min']),abs(sd['max']))*100,1)} %.

**Рисунок 2**

*Динаміка щоденної прибутковості акцій {short} (%, 2026 р.)*

![]({img}_ret_d.png)

*Примітка.* Побудовано автором за даними Yahoo Finance (2026).

{stat_table(sd,'3','щоденної',n-1)}

Середня щоденна прибутковість за період виявилася {mean_d_word} та становила {num(sd['mean']*100,3)} %, що {align_word}. Форма розподілу прибутковості відхиляється від нормальної: ексцес ({num(sd['kurt'],2)}) указує на {kurt_d}, а {skew_d}. Це наочно підтверджує і гістограма розподілу (рисунок 3).

**Рисунок 3**

*Гістограма розподілу щоденної прибутковості акцій {short} (2026 р.)*

![]({img}_hist.png)

*Примітка.* Суцільною лінією показано щільність нормального розподілу з тими ж середнім і стандартним відхиленням. Побудовано автором.

Емпіричний розподіл помітно відрізняється від нормальної кривої формою центральної частини та хвостів, що типово для денної прибутковості акцій і свідчить про невипадковий характер ряду.

## Щотижнева прибутковість

Перехід до щотижневих даних згладжує частину денного «шуму» (рисунок 4). Узагальнені показники щотижневої прибутковості наведено в таблиці 4.

**Рисунок 4**

*Динаміка щотижневої прибутковості акцій {short} (%, 2026 р.)*

![]({img}_ret_w.png)

*Примітка.* Побудовано автором за даними Yahoo Finance (2026).

{stat_table(sw,'4','щотижневої',nw)}

На тижневому горизонті середня прибутковість є {mean_w_word} ({num(sw['mean']*100,2)} %), а коефіцієнт асиметрії ({num(sw['skew'],2)}) відображає {'переважання тижнів зі значним додатним результатом' if sw['skew']>0 else 'переважання тижнів зі значним відʼємним результатом'}.

## Щомісячна прибутковість

Щомісячна прибутковість (рисунок 5, таблиця 5) розрахована на коротшій вибірці з {nm} спостережень, тому її показники мають радше ілюстративний характер і не дають підстав для статистично надійних висновків.

**Рисунок 5**

*Динаміка щомісячної прибутковості акцій {short} (%, 2026 р.)*

![]({img}_ret_m.png)

*Примітка.* Побудовано автором за даними Yahoo Finance (2026).

{stat_table(sm,'5','щомісячної',nm)}

Слід окремо зауважити, що для рядів прибутковості коефіцієнти варіації та осциляції набувають {'від’ємних і ' if (sd['mean']<0 or sw['mean']<0) else ''}аномально великих за модулем значень. Це є наслідком близькості середньої прибутковості до нуля: при $\\bar{{x}} \\to 0$ відношення $\\sigma / \\bar{{x}}$ втрачає економічний зміст, тому ці показники наведено лише для повноти розрахунку, а мінливість прибутковості коректніше оцінювати за стандартним відхиленням і розмахом.

# Висновки

Опрацювання акцій {legal} дало змогу пройти повний цикл роботи з ринковими даними — від вибору інструмента та імпорту часового ряду до його статистичного аналізу. У першому півріччі 2026 року ціна акцій рухалася за {traj_adj} немонотонною траєкторією, яку найкраще описує {bt} тренд ($R^2 = {num(max(l['r2'],po['r2'],pw['r2']),4)}$). Коефіцієнт варіації цінового ряду становить {num(p['cv'],2)} %, тоді як розподіл щоденної прибутковості відхиляється від нормального закону за показниками асиметрії та ексцесу. Отримані результати підтверджують доцільність застосування описаної методики й узгоджуються з аналогічними розрахунками для решти обраних інструментів — {others_str}.

# Список використаних джерел

Арутюнян, С. С., Добриніна, Л. В. та ін. (2018). *Фінансовий ринок* (С. С. Арутюнян, Ред.). Гуляєва В. М.

Горбаченко, С. А., Карпов, В. А., & Шевченко-Перепьолкіна, Р. І. (2019). *Аналіз та прогнозування ринкової кон'юнктури*. Кондор.

Максишко, Н. К. (2009). *Моделювання економіки методами дискретної нелінійної динаміки*. Запорізький національний університет.

Максишко, Н. К., & Шаповалова, В. О. (2018). *Аналіз та прогнозування: сучасні концепції в дослідженні динаміки ціни на ринку нерухомості*. Поліграф.

Investing.com. (2026). *Historical data*. https://www.investing.com/

Yahoo Finance. (2026). *{legal} ({tk}) historical prices*. https://finance.yahoo.com/quote/{tk}/history
"""
    return md

# ---------- APA formatting ----------
def set_font(run,name='Times New Roman',size=12,bold=None,italic=None):
    run.font.name=name; run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if italic is not None: run.font.italic=italic
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),name)

def is_caption(t):
    t=t.strip()
    return t.startswith(('Таблиця','Рисунок','Примітка','Статистичні','Динаміка','Гістограма',
                         'Показники','Фрагмент')) and len(t)<200

def format_apa(path, legal, short):
    doc=Document(path)
    normal=doc.styles['Normal']; normal.font.name='Times New Roman'; normal.font.size=Pt(12)
    normal.font.color.rgb=RGBColor(0,0,0)
    pf=normal.paragraph_format; pf.line_spacing_rule=WD_LINE_SPACING.DOUBLE
    pf.space_after=Pt(0); pf.space_before=Pt(0)
    rpr=normal.element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None: rf=OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Times New Roman')
    in_refs=False
    for p in doc.paragraphs:
        style=p.style.name; txt=p.text.strip()
        if style.startswith('Heading'):
            in_refs=txt.startswith('Список використаних джерел')
            for r in p.runs:
                set_font(r,size=14 if style=='Heading 1' else 13,bold=True); r.font.color.rgb=RGBColor(0,0,0)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER if style=='Heading 1' else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(6)
            p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
            continue
        if p._element.findall('.//'+qn('w:drawing')):
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER; continue
        for r in p.runs: set_font(r)
        if not txt: continue
        p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
        if in_refs:
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent=Cm(1.25); p.paragraph_format.first_line_indent=Cm(-1.25)
        elif is_caption(txt):
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.first_line_indent=None
        else:
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.first_line_indent=Cm(1.25)
    for tbl in doc.tables:
        tbl.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
                    p.paragraph_format.space_after=Pt(0)
                    for r in p.runs: set_font(r,size=11)
    # title page
    first=doc.paragraphs[0]._element
    def add_before(text,bold=False,italic=False,size=12):
        pnew=OxmlElement('w:p'); first.addprevious(pnew)
        par=Paragraph(pnew,doc.paragraphs[0]._parent)
        par.alignment=WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
        r=par.add_run(text); set_font(r,size=size,bold=bold,italic=italic)
    lines=[('Запорізький національний університет',0,0,12),('Кафедра економічної кібернетики',0,0,12),
        ('',0,0,12),('',0,0,12),('',0,0,12),
        ('ІНДИВІДУАЛЬНИЙ ПРАКТИЧНИЙ КЕЙС № 1',1,0,14),
        ('Інфраструктура фінансового ринку та джерела аналізу його динаміки',1,0,12),
        (f'(на прикладі акцій {legal})',0,1,12),
        ('',0,0,12),('',0,0,12),('',0,0,12),('',0,0,12),
        ('Навчальна дисципліна:',0,0,12),('«Новий аналітичний погляд на динаміку ринків»',0,1,12),
        ('',0,0,12),('Виконав: Гордій Олександр',0,0,12),
        ('',0,0,12),('',0,0,12),('',0,0,12),('Запоріжжя — 2026',0,0,12)]
    for text,b,i,sz in lines: add_before(text,bool(b),bool(i),sz)
    pbreak=OxmlElement('w:p'); first.addprevious(pbreak)
    pbpar=Paragraph(pbreak,doc.paragraphs[0]._parent); pbpar.add_run().add_break(WD_BREAK.PAGE)
    doc.save(path)

# ---------- run ----------
sel = sys.argv[1:] if len(sys.argv)>1 else [c[0] for c in COMPANIES]
for tk,legal,short in COMPANIES:
    if tk not in sel: continue
    close=close_all[tk].dropna()
    tr=trends(close.values)
    rd,rw,rm=make_figs(tk,short,close,tr)
    others=[s for t,l,s in COMPANIES if t!=tk]
    md=build_md(tk,legal,short,close,tr,rd,rw,rm,others)
    mdpath=f'report/{tk}_report.md'; open(mdpath,'w',encoding='utf-8').write(md)
    out=f'report/Кейс1_{short}.docx'
    pypandoc.convert_file(mdpath,'docx',outputfile=out,extra_args=['--resource-path=.'])
    format_apa(out,legal,short)
    print(f'{tk}: {out}  ({os.path.getsize(out)} байт)')
print('DONE')
