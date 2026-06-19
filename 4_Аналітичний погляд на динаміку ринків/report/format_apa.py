import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = 'report/Кейс1_Microsoft.docx'
doc = Document(SRC)

def set_font(run, name='Times New Roman', size=12, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii','w:hAnsi','w:cs'):
        rfonts.set(qn(a), name)

# --- base style ---
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor(0,0,0)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0); pf.space_before = Pt(0)
rpr = normal.element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
if rf is None: rf = OxmlElement('w:rFonts'); rpr.append(rf)
for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Times New Roman')

def is_caption(txt):
    t = txt.strip()
    return t.startswith(('Таблиця','Рисунок','Примітка','Статистичні','Динаміка','Гістограма',
                         'Показники','Фрагмент')) and len(t) < 200

# --- body paragraphs ---
in_refs = False
for p in doc.paragraphs:
    style = p.style.name
    txt = p.text.strip()
    # headings
    if style.startswith('Heading'):
        in_refs = txt.startswith('Список використаних джерел')
        for r in p.runs:
            set_font(r, size=14 if style=='Heading 1' else 13, bold=True)
            r.font.color.rgb = RGBColor(0,0,0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if style=='Heading 1' else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        continue
    # figure (image) paragraph -> center
    if p._element.findall('.//'+qn('w:drawing')):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue
    for r in p.runs:
        set_font(r)
    if not txt:
        continue
    # table number / caption lines: center, no indent
    if txt in ('','…') or txt.startswith(('**',)):
        pass
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if in_refs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
    elif is_caption(txt) or txt.startswith('Примітка'):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = None
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(1.25)

# --- tables: font + centering ---
for tbl in doc.tables:
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=11)

# --- build APA title page, prepend ---
def title_para(text, bold=False, italic=False, size=12, before=0, after=0):
    pnew = OxmlElement('w:p')
    doc.paragraphs[0]._element.addprevious(pnew)
    from docx.text.paragraph import Paragraph
    par = Paragraph(pnew, doc.paragraphs[0]._parent)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    par.paragraph_format.space_before = Pt(before); par.paragraph_format.space_after = Pt(after)
    r = par.add_run(text); set_font(r, size=size, bold=bold, italic=italic)
    return par

# prepend in reverse visual order is tricky; build list then insert before first body para
first = doc.paragraphs[0]._element
def add_before(first, text, bold=False, italic=False, size=12, blank_after=0):
    from docx.text.paragraph import Paragraph
    pnew = OxmlElement('w:p'); first.addprevious(pnew)
    par = Paragraph(pnew, doc.paragraphs[0]._parent)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    r = par.add_run(text); set_font(r, size=size, bold=bold, italic=italic)
    return par

title_lines = [
    ('Запорізький національний університет', False, False, 12),
    ('Кафедра економічної кібернетики', False, False, 12),
    ('', False, False, 12),
    ('', False, False, 12),
    ('', False, False, 12),
    ('ІНДИВІДУАЛЬНИЙ ПРАКТИЧНИЙ КЕЙС № 1', True, False, 14),
    ('Інфраструктура фінансового ринку та джерела аналізу його динаміки', True, False, 12),
    ('(на прикладі акцій Microsoft Corporation)', False, True, 12),
    ('', False, False, 12),
    ('', False, False, 12),
    ('', False, False, 12),
    ('', False, False, 12),
    ('Навчальна дисципліна:', False, False, 12),
    ('«Новий аналітичний погляд на динаміку ринків»', False, True, 12),
    ('', False, False, 12),
    ('Виконав: Гордій Олександр', False, False, 12),
    ('', False, False, 12),
    ('', False, False, 12),
    ('', False, False, 12),
    ('Запоріжжя — 2026', False, False, 12),
]
for text, b, i, sz in title_lines:
    add_before(first, text, bold=b, italic=i, size=sz)

# page break after title page (before first heading)
from docx.text.paragraph import Paragraph
pbreak = OxmlElement('w:p'); first.addprevious(pbreak)
pbpar = Paragraph(pbreak, doc.paragraphs[0]._parent)
run = pbpar.add_run(); run.add_break(WD_BREAK.PAGE)

doc.save(SRC)
print('APA formatting applied ->', SRC)
print('Параграфів:', len(doc.paragraphs), '| Таблиць:', len(doc.tables))
