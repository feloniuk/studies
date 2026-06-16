import sys, os, glob
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

base = 'D:\\feloniuk\\studies\\2_\u0422\u0430\u0439\u043c \u043c\u0435\u043d\u0435\u0434\u0436\u043c\u0435\u043d\u0442\\' + '\u041c\u043e\u0434\u0443\u043b\u044c 3'

# Find files using glob
for f in os.listdir(base):
    full = os.path.join(base, f)
    if '\u0410\u043d\u0430\u043b\u0456\u0437' in f:
        print(f"Reading: {f}")
        doc = Document(full)
        for p in doc.paragraphs:
            print(p.text)
        for i, table in enumerate(doc.tables):
            print(f"\n--- TABLE {i+1} ---")
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                print(' | '.join(cells))
